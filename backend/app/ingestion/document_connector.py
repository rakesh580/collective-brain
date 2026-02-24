import os
import re
from datetime import datetime
from pathlib import Path

from app.ingestion.base import BaseConnector, ParsedChunk
from app.ingestion.chunker import ChunkingStrategy


class DocumentConnector(BaseConnector):
    """Connector for PDF, DOCX, and TXT files."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.chunker = ChunkingStrategy(chunk_size, chunk_overlap)

    def source_type(self) -> str:
        return "document"

    def parse(self, source_input: str) -> list[ParsedChunk]:
        """Parse a single document file (PDF, DOCX, or TXT)."""
        file_path = Path(source_input)
        ext = file_path.suffix.lower()

        if ext == ".pdf":
            text = self._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text = self._extract_docx(file_path)
        elif ext == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        if not text.strip():
            return []

        filename = file_path.name
        title = file_path.stem.replace("-", " ").replace("_", " ").title()
        topics = self._extract_topics(filename, text)

        # Chunk the extracted text
        sub_chunks = self.chunker.split(
            f"[Document: {title}]\n{text}",
            metadata={"source_ref": filename, "file_type": ext},
        )

        chunks = []
        for sc in sub_chunks:
            chunks.append(
                ParsedChunk(
                    text=sc["text"],
                    source_type="document",
                    source_ref=filename,
                    author=None,
                    timestamp=datetime.fromtimestamp(os.path.getmtime(file_path)),
                    topics=topics,
                    chunk_metadata=sc["metadata"],
                )
            )
        return chunks

    def extract_members(self, source_input: str) -> list[dict]:
        """Documents don't have structured member info."""
        return []

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF using PyMuPDF (fitz)."""
        import fitz  # pymupdf

        text_parts = []
        with fitz.open(str(file_path)) as doc:
            for page in doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from a DOCX file using python-docx."""
        from docx import Document

        doc = Document(str(file_path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)

    def _extract_topics(self, filename: str, text: str) -> list[str]:
        """Extract topics from the filename and first few lines of text."""
        topics = set()

        # From filename
        stem = Path(filename).stem.replace("-", " ").replace("_", " ").lower()
        words = [w for w in stem.split() if len(w) > 2]
        topics.update(words[:3])

        # From first heading-like lines
        for line in text.split("\n")[:20]:
            line = line.strip()
            if line and len(line) < 80 and not line.endswith("."):
                cleaned = re.sub(r"[^a-zA-Z0-9\s]", "", line).lower().strip()
                if cleaned and len(cleaned.split()) <= 5:
                    topics.add(cleaned)
                    if len(topics) >= 5:
                        break

        return list(topics)[:5]
